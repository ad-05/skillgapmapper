"""
Resume Parser & Skill Extractor
================================
Extracts text from PDF/DOCX resumes, detects skills using context-aware
NLP pattern matching, and estimates proficiency levels from experience signals.

No external NLP library required — pure Python regex + linguistic rules.
"""

import re
import io
import json
import math
from typing import Optional

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""

    # ── Attempt 1: pypdf ──────────────────────────────────────────────────────
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages  = [page.extract_text() or "" for page in reader.pages]
        text   = "\n".join(pages).strip()
    except Exception:
        pass

    # ── Attempt 2: pdfplumber (handles complex layouts better) ───────────────
    if len(text) < 100:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(pages).strip()
        except Exception as e:
            if not text:
                raise ValueError(f"Could not read PDF: {e}")

    if not text:
        raise ValueError(
            "Could not extract text from this PDF. "
            "It may be a scanned image — try a text-based PDF or paste your resume as a .txt file."
        )

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc  = Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text
        return text
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload PDF, DOCX, or TXT.")


# ── Skill keyword mapping ─────────────────────────────────────────────────────
# Each skill maps to a list of keywords/phrases that indicate its presence.
# Longer, more specific phrases are weighted higher.

SKILL_KEYWORDS = {
    "Python": [
        "python", "django", "flask", "fastapi", "pandas", "numpy", "scipy",
        "pytorch", "tensorflow", "keras", "sklearn", "scikit-learn", "jupyter",
        "anaconda", "pycharm", "pip", "virtualenv", "asyncio", "celery",
    ],
    "JavaScript": [
        "javascript", "js", "es6", "es2015", "ecmascript", "jquery",
        "vanilla js", "node.js", "nodejs", "npm", "yarn", "webpack",
        "babel", "eslint", "jest", "mocha",
    ],
    "React": [
        "react", "reactjs", "react.js", "jsx", "redux", "react hooks",
        "usestate", "useeffect", "react router", "next.js", "nextjs",
        "gatsby", "react native",
    ],
    "Node.js": [
        "node.js", "nodejs", "node", "express", "expressjs", "koa",
        "nestjs", "socket.io", "npm", "yarn",
    ],
    "SQL": [
        "sql", "mysql", "postgresql", "postgres", "sqlite", "oracle",
        "mssql", "sql server", "t-sql", "pl/sql", "stored procedures",
        "query optimization", "joins", "indexing", "normalization",
    ],
    "REST APIs": [
        "rest", "restful", "api", "rest api", "http", "json api",
        "openapi", "swagger", "postman", "api design", "microservices",
        "web services", "endpoint", "graphql",
    ],
    "Git": [
        "git", "github", "gitlab", "bitbucket", "version control",
        "git flow", "pull request", "code review", "branching",
        "merge", "rebase", "ci/cd",
    ],
    "CSS": [
        "css", "css3", "sass", "scss", "less", "tailwind", "bootstrap",
        "material ui", "styled components", "flexbox", "css grid",
        "responsive design", "media queries",
    ],
    "TypeScript": [
        "typescript", "ts", "type-safe", "angular",
    ],
    "Docker": [
        "docker", "dockerfile", "container", "containerization",
        "docker-compose", "docker swarm", "image", "registry",
    ],
    "Machine Learning": [
        "machine learning", "ml", "supervised learning", "unsupervised learning",
        "classification", "regression", "clustering", "scikit-learn",
        "feature engineering", "model training", "cross-validation",
        "random forest", "svm", "gradient boosting", "xgboost",
    ],
    "Statistics": [
        "statistics", "statistical", "hypothesis testing", "probability",
        "regression analysis", "anova", "chi-square", "confidence interval",
        "p-value", "bayesian", "a/b testing", "statistical modeling",
    ],
    "Data Visualization": [
        "data visualization", "tableau", "power bi", "matplotlib",
        "seaborn", "plotly", "d3.js", "grafana", "kibana",
        "dashboard", "charts", "reporting",
    ],
    "Pandas": [
        "pandas", "dataframe", "data manipulation", "data wrangling",
        "data cleaning", "data preprocessing",
    ],
    "Deep Learning": [
        "deep learning", "neural network", "cnn", "rnn", "lstm",
        "transformer", "bert", "gpt", "attention mechanism",
        "backpropagation", "gpu training", "pytorch", "tensorflow", "keras",
    ],
    "Feature Engineering": [
        "feature engineering", "feature selection", "feature extraction",
        "dimensionality reduction", "pca", "one-hot encoding",
        "target encoding", "feature importance",
    ],
    "Big Data": [
        "big data", "hadoop", "spark", "apache spark", "pyspark",
        "hive", "kafka", "flink", "data lake", "hdfs", "mapreduce",
        "databricks",
    ],
    "Kubernetes": [
        "kubernetes", "k8s", "kubectl", "helm", "pod", "deployment",
        "service mesh", "istio", "eks", "aks", "gke",
    ],
    "Linux": [
        "linux", "ubuntu", "centos", "debian", "unix", "bash",
        "shell scripting", "shell script", "command line", "terminal",
        "ssh", "systemd", "cron",
    ],
    "CI/CD": [
        "ci/cd", "continuous integration", "continuous deployment",
        "jenkins", "github actions", "gitlab ci", "travis ci",
        "circle ci", "pipeline", "automated testing", "devops",
    ],
    "AWS": [
        "aws", "amazon web services", "ec2", "s3", "lambda", "rds",
        "cloudformation", "iam", "vpc", "route53", "cloudwatch",
        "elastic beanstalk", "ecs", "eks", "dynamodb",
    ],
    "Terraform": [
        "terraform", "infrastructure as code", "iac", "hcl",
        "terraform modules", "state management",
    ],
    "Networking": [
        "networking", "tcp/ip", "dns", "http", "https", "ssl", "tls",
        "firewall", "vpn", "load balancer", "network security",
        "vlan", "subnets", "routing",
    ],
    "Monitoring": [
        "monitoring", "prometheus", "grafana", "datadog", "newrelic",
        "splunk", "elk stack", "alerting", "slo", "sli", "observability",
        "logging", "tracing", "apm",
    ],
    "Ethical Hacking": [
        "ethical hacking", "penetration testing", "pen test", "pentest",
        "vulnerability assessment", "burp suite", "metasploit", "nmap",
        "kali linux", "ceh", "oscp", "red team",
    ],
    "Cryptography": [
        "cryptography", "encryption", "decryption", "aes", "rsa",
        "sha", "hashing", "ssl", "tls", "pki", "certificate",
        "digital signature", "public key",
    ],
    "Incident Response": [
        "incident response", "incident management", "security incident",
        "forensics", "digital forensics", "post-mortem", "runbook",
        "playbook", "soc", "triage",
    ],
    "SIEM Tools": [
        "siem", "splunk", "qradar", "arcsight", "sentinel",
        "log management", "security operations", "soc",
    ],
    "Risk Assessment": [
        "risk assessment", "risk management", "risk analysis",
        "threat modeling", "vulnerability management", "cvss",
        "iso 27001", "nist", "compliance audit",
    ],
    "Firewalls": [
        "firewall", "palo alto", "cisco asa", "fortinet",
        "network security", "acl", "packet filtering",
    ],
    "Compliance": [
        "compliance", "gdpr", "hipaa", "pci dss", "iso 27001",
        "soc 2", "regulatory", "audit", "governance",
    ],
    "TensorFlow/PyTorch": [
        "tensorflow", "pytorch", "keras", "neural network",
        "deep learning", "model deployment", "onnx",
    ],
    "Mathematics": [
        "mathematics", "linear algebra", "calculus", "differential equations",
        "numerical methods", "optimization", "matrix", "eigenvalue",
        "fourier", "probability theory",
    ],
    "Data Processing": [
        "data processing", "etl", "data pipeline", "data engineering",
        "data ingestion", "data transformation", "airflow", "luigi",
    ],
    "MLOps": [
        "mlops", "model deployment", "model serving", "mlflow",
        "kubeflow", "sagemaker", "model monitoring", "model registry",
        "feature store", "a/b testing models",
    ],
    "Cloud Platforms": [
        "cloud", "aws", "azure", "gcp", "google cloud", "cloud computing",
        "serverless", "paas", "iaas", "saas",
    ],
    "Product Strategy": [
        "product strategy", "product vision", "go-to-market", "product roadmap",
        "okr", "kpi", "north star", "product-market fit",
    ],
    "Agile/Scrum": [
        "agile", "scrum", "kanban", "sprint", "backlog", "standup",
        "retrospective", "user story", "velocity", "jira", "confluence",
        "product owner",
    ],
    "Data Analysis": [
        "data analysis", "data analytics", "business intelligence",
        "excel", "google sheets", "pivot table", "sql analytics",
        "looker", "power bi", "tableau",
    ],
    "Communication": [
        "communication", "presentation", "stakeholder", "cross-functional",
        "leadership", "team lead", "mentoring", "public speaking",
        "documentation", "technical writing",
    ],
    "User Research": [
        "user research", "ux research", "usability testing", "user interview",
        "persona", "user journey", "a/b testing", "survey",
    ],
    "Roadmapping": [
        "roadmap", "product roadmap", "release planning", "feature prioritization",
        "strategic planning", "quarterly planning",
    ],
    "Stakeholder Management": [
        "stakeholder management", "stakeholder", "executive communication",
        "cross-functional", "negotiation", "alignment",
    ],
    "Market Research": [
        "market research", "competitive analysis", "market analysis",
        "tam", "sam", "som", "industry analysis", "customer discovery",
    ],
    "Leadership": [
        "leadership", "team lead", "tech lead", "engineering manager",
        "mentoring", "coaching", "people management", "director",
    ],
}

# ── Experience signals — these boost proficiency estimates ────────────────────
# Patterns that indicate YEARS of experience with a skill
YEAR_PATTERNS = [
    r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience\s+(?:with\s+|in\s+)?)?({skill})',
    r'({skill})\s+(?:for\s+)?(\d+)\+?\s*(?:years?|yrs?)',
    r'(\d+)\+?\s*(?:years?|yrs?)\s+({skill})',
]

# Signals that indicate senior/expert level
EXPERT_SIGNALS = [
    "led", "lead", "architected", "architect", "designed", "built from scratch",
    "created", "developed", "published", "authored", "contributed to",
    "open source", "production", "scaled", "senior", "principal",
    "expert", "advanced", "deep knowledge", "extensive experience",
    "5+ years", "6+ years", "7+ years", "8+ years",
]

# Signals that indicate intermediate level
INTERMEDIATE_SIGNALS = [
    "developed", "implemented", "worked with", "used", "applied",
    "experience with", "familiar with", "proficient", "comfortable",
    "2 years", "3 years", "4 years", "hands-on",
]

# Signals that indicate beginner level
BEGINNER_SIGNALS = [
    "learning", "studying", "beginner", "basic", "introductory",
    "coursework", "academic", "college project", "course",
    "exposure to", "awareness of", "recently started",
]


def _extract_years(text: str, skill_keywords: list[str]) -> Optional[int]:
    """Try to extract explicit year counts for a skill from text."""
    text_lower = text.lower()
    for kw in skill_keywords[:3]:  # check top keywords only
        kw_escaped = re.escape(kw.lower())
        patterns = [
            rf'(\d+)\+?\s*(?:years?|yrs?)[\w\s]*?{kw_escaped}',
            rf'{kw_escaped}[\w\s]*?(\d+)\+?\s*(?:years?|yrs?)',
        ]
        for pat in patterns:
            m = re.search(pat, text_lower)
            if m:
                return int(m.group(1))
    return None


def _count_occurrences(text: str, keywords: list[str]) -> int:
    """Count how many times any keyword appears in text."""
    text_lower = text.lower()
    count = 0
    for kw in keywords:
        # Use word boundary matching where possible
        pattern = rf'\b{re.escape(kw.lower())}\b'
        count += len(re.findall(pattern, text_lower))
    return count


def _get_context_window(text: str, keyword: str, window: int = 150) -> list[str]:
    """Get text windows around each occurrence of a keyword."""
    text_lower = text.lower()
    kw_lower   = keyword.lower()
    windows    = []
    start      = 0
    while True:
        idx = text_lower.find(kw_lower, start)
        if idx == -1:
            break
        w_start = max(0, idx - window)
        w_end   = min(len(text), idx + len(keyword) + window)
        windows.append(text[w_start:w_end].lower())
        start = idx + 1
    return windows


def estimate_proficiency(text: str, skill: str, keywords: list[str]) -> Optional[int]:
    """
    Estimate proficiency level (0, 3, 6, 8, 10) for a skill
    based on keyword presence and context signals.
    Returns None if skill not detected at all.
    """
    text_lower = text.lower()

    # Check if ANY keyword is present
    occurrences = _count_occurrences(text, keywords)
    if occurrences == 0:
        return None

    # Get all context windows around skill mentions
    all_contexts = []
    for kw in keywords[:5]:
        all_contexts.extend(_get_context_window(text, kw))
    combined_context = " ".join(all_contexts)

    # Check for explicit year mentions
    years = _extract_years(text, keywords)
    if years is not None:
        if years >= 5:   return 10
        if years >= 3:   return 8
        if years >= 2:   return 6
        if years >= 1:   return 3
        return 3

    # Score based on context signals
    expert_score       = sum(1 for s in EXPERT_SIGNALS if s in combined_context)
    intermediate_score = sum(1 for s in INTERMEDIATE_SIGNALS if s in combined_context)
    beginner_score     = sum(1 for s in BEGINNER_SIGNALS if s in combined_context)

    # Frequency-based score (more mentions = more experience)
    freq_score = min(occurrences / 3.0, 3.0)

    # Decision logic
    if beginner_score > 0 and expert_score == 0:
        return 3   # Beginner

    total_positive = expert_score + intermediate_score + freq_score

    if expert_score >= 2 or total_positive >= 5:
        return 10  # Expert
    if expert_score >= 1 or total_positive >= 3:
        return 8   # Advanced
    if intermediate_score >= 1 or total_positive >= 1.5:
        return 6   # Intermediate
    if occurrences >= 1:
        return 3   # Beginner (mentioned but no strong context)

    return None


# ── Section detection ─────────────────────────────────────────────────────────

SECTION_HEADERS = {
    "skills":      r'\b(skills?|technical skills?|core competencies|technologies|expertise|proficiencies)\b',
    "experience":  r'\b(experience|work experience|employment|work history|professional experience|career)\b',
    "education":   r'\b(education|academic|qualifications|degrees?|university|college)\b',
    "projects":    r'\b(projects?|personal projects?|side projects?|portfolio)\b',
    "certifications": r'\b(certifications?|certificates?|licenses?|credentials?)\b',
    "summary":     r'\b(summary|profile|objective|about|overview)\b',
}


def detect_sections(text: str) -> dict:
    """Split resume text into sections."""
    lines    = text.split("\n")
    sections = {"full": text, "skills": "", "experience": "", "projects": ""}
    current  = "full"

    for line in lines:
        line_lower = line.lower().strip()
        detected   = False
        for sec, pattern in SECTION_HEADERS.items():
            if re.search(pattern, line_lower) and len(line_lower) < 50:
                current   = sec
                detected  = True
                break
        if current in sections:
            sections[current] += " " + line

    return sections


# ── Personal info extraction ──────────────────────────────────────────────────

def extract_personal_info(text: str) -> dict:
    """Extract name, email, phone from resume text."""
    info = {}

    # Email
    email_match = re.search(r'[\w\.\+\-]+@[\w\.\-]+\.\w{2,}', text)
    if email_match:
        info["email"] = email_match.group()

    # Phone
    phone_match = re.search(r'[\+\(]?[\d\s\-\(\)]{10,15}', text)
    if phone_match:
        cleaned = re.sub(r'[\s\-\(\)]', '', phone_match.group())
        if len(cleaned) >= 10:
            info["phone"] = phone_match.group().strip()

    # LinkedIn
    linkedin = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    if linkedin:
        info["linkedin"] = "https://" + linkedin.group()

    # GitHub
    github = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
    if github:
        info["github"] = "https://" + github.group()

    # Years of experience (total)
    exp_match = re.search(r'(\d+)\+?\s*years?\s+(?:of\s+)?(?:overall\s+)?experience',
                          text, re.IGNORECASE)
    if exp_match:
        info["total_years"] = int(exp_match.group(1))

    return info


# ── Main parse function ───────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Full resume parsing pipeline.

    Returns:
        {
          "extracted_skills": { skill_name: proficiency_level (0-10) },
          "detected_count": int,
          "personal_info": { email, phone, ... },
          "word_count": int,
          "sections_found": [list of detected section names],
          "skill_evidence": { skill: [list of context snippets] },
          "confidence_scores": { skill: confidence (0-1) },
          "raw_text_preview": first 500 chars,
        }
    """
    # Step 1: Extract text
    raw_text = extract_text(file_bytes, filename)
    if len(raw_text.strip()) < 20:
        raise ValueError("Could not extract meaningful text from the file. "
                         "Try a text-based PDF rather than a scanned image.")

    # Step 2: Detect sections
    sections = detect_sections(raw_text)

    # Step 3: Extract personal info
    personal_info = extract_personal_info(raw_text)

    # Step 4: Score every skill
    extracted_skills   = {}
    skill_evidence     = {}
    confidence_scores  = {}

    for skill, keywords in SKILL_KEYWORDS.items():
        # Give higher weight to skills section + experience section
        skill_text = (
            sections.get("skills", "")     * 3 +   # skills section x3
            sections.get("experience", "") * 2 +   # experience x2
            sections.get("projects", "")   * 2 +   # projects x2
            raw_text                                # full text x1
        )

        level = estimate_proficiency(skill_text, skill, keywords)

        if level is not None and level > 0:
            extracted_skills[skill]  = level
            occurrences = _count_occurrences(raw_text, keywords)
            # Confidence: based on number of distinct keyword matches
            unique_kws_found = sum(1 for kw in keywords
                                   if kw.lower() in raw_text.lower())
            confidence_scores[skill] = min(unique_kws_found / max(len(keywords) * 0.3, 1), 1.0)
            confidence_scores[skill] = round(confidence_scores[skill], 2)

            # Collect evidence snippets
            snippets = []
            for kw in keywords[:3]:
                windows = _get_context_window(raw_text, kw, window=80)
                snippets.extend(windows[:2])
            skill_evidence[skill] = list(set(s.strip() for s in snippets[:3]))

    # Step 5: Detect sections actually found
    sections_found = [
        sec for sec, content in sections.items()
        if sec != "full" and len(content.strip()) > 20
    ]

    return {
        "extracted_skills":   extracted_skills,
        "detected_count":     len(extracted_skills),
        "personal_info":      personal_info,
        "word_count":         len(raw_text.split()),
        "sections_found":     sections_found,
        "skill_evidence":     skill_evidence,
        "confidence_scores":  confidence_scores,
        "raw_text_preview":   raw_text[:500].strip(),
    }


# ── Resume-role fit scoring ───────────────────────────────────────────────────

def compute_resume_fit(
    extracted_skills: dict,
    role_skills: dict,
    role_name: str,
) -> dict:
    """
    Compute how well a resume fits a specific role.
    Returns a detailed fit report separate from the assessment score.
    """
    required = {k: v for k, v in role_skills.items() if v > 0}
    if not required:
        return {}

    # Coverage: what % of required skills appear in the resume at all
    covered_skills = {s for s in required if extracted_skills.get(s, 0) > 0}
    coverage_pct   = len(covered_skills) / len(required) * 100

    # Strength match: for covered skills, how close is the level
    level_scores = []
    for skill, req_level in required.items():
        user_level = extracted_skills.get(skill, 0)
        if user_level > 0:
            # Score = min(user/required, 1.0)
            level_scores.append(min(user_level / req_level, 1.0))

    avg_level_match = sum(level_scores) / len(level_scores) * 100 if level_scores else 0

    # Combined fit score
    fit_score = (coverage_pct * 0.5 + avg_level_match * 0.5)

    # Missing critical skills (required >= 7)
    critical_missing = [
        s for s, v in required.items()
        if v >= 7 and extracted_skills.get(s, 0) == 0
    ]

    # Strong matches (user meets or exceeds requirement)
    strong_matches = [
        s for s in covered_skills
        if extracted_skills.get(s, 0) >= required[s]
    ]

    # Weak matches (present but below requirement)
    weak_matches = [
        {"skill": s, "resume_level": extracted_skills[s], "required": required[s]}
        for s in covered_skills
        if extracted_skills.get(s, 0) < required[s]
    ]
    weak_matches.sort(key=lambda x: x["required"] - x["resume_level"], reverse=True)

    fit_label = (
        "Excellent Fit 🎯"  if fit_score >= 80 else
        "Good Fit ✅"       if fit_score >= 65 else
        "Moderate Fit 💪"   if fit_score >= 45 else
        "Needs Work 📚"     if fit_score >= 25 else
        "Low Fit 🌱"
    )

    return {
        "role_name":        role_name,
        "fit_score":        round(fit_score, 1),
        "fit_label":        fit_label,
        "coverage_pct":     round(coverage_pct, 1),
        "avg_level_match":  round(avg_level_match, 1),
        "required_count":   len(required),
        "covered_count":    len(covered_skills),
        "critical_missing": critical_missing,
        "strong_matches":   strong_matches,
        "weak_matches":     weak_matches[:6],
    }